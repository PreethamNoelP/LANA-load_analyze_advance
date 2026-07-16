import io

import pandas as pd
from fpdf import FPDF
from fpdf.enums import XPos, YPos
from docx import Document


def _pdf_text(text: str) -> str:
    """Make text safe for fpdf2's built-in core fonts (latin-1 only).

    Characters outside latin-1 are replaced with '?' instead of crashing
    the export. Swap the core font for a Unicode TTF to lift this limit.
    """
    return str(text).encode("latin-1", "replace").decode("latin-1")


def generate_pdf_report(df: pd.DataFrame, context: str) -> bytes:
    """Generate a PDF summary report for the uploaded dataset."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("helvetica", "B", 18)
    pdf.cell(0, 12, "LANA Analysis Report", align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)

    pdf.set_font("helvetica", "B", 13)
    pdf.cell(0, 9, "Dataset Overview", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("helvetica", "", 10)
    pdf.cell(0, 7, f"  Rows: {len(df):,}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.cell(0, 7, f"  Columns: {len(df.columns)}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.multi_cell(0, 7, _pdf_text(f"  Column names: {', '.join(df.columns.tolist())}"),
                   new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(3)

    pdf.set_font("helvetica", "B", 13)
    pdf.cell(0, 9, "Data Profile", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("helvetica", "", 9)
    for line in context.splitlines():
        pdf.multi_cell(0, 5, _pdf_text(f"  {line}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(3)

    numeric_cols = df.select_dtypes("number").columns.tolist()
    if numeric_cols:
        pdf.set_font("helvetica", "B", 13)
        pdf.cell(0, 9, "Numeric Column Summary", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("helvetica", "", 9)
        for col in numeric_cols:
            s = df[col].dropna()
            pdf.multi_cell(
                0, 5,
                _pdf_text(
                    f"  {col}: "
                    f"mean={s.mean():.4g}, std={s.std():.4g}, "
                    f"min={s.min():.4g}, max={s.max():.4g}, "
                    f"nulls={df[col].isnull().sum()}"
                ),
                new_x=XPos.LMARGIN, new_y=YPos.NEXT,
            )
        pdf.ln(3)

    return bytes(pdf.output())


def generate_word_report(df: pd.DataFrame, context: str) -> bytes:
    """Generate a Word document summary report for the uploaded dataset."""
    doc = Document()
    doc.add_heading("LANA Analysis Report", level=0)

    doc.add_heading("Dataset Overview", level=1)
    doc.add_paragraph(f"Rows: {len(df):,}", style="List Bullet")
    doc.add_paragraph(f"Columns: {len(df.columns)}", style="List Bullet")
    doc.add_paragraph(f"Column names: {', '.join(df.columns.tolist())}", style="List Bullet")

    doc.add_heading("Data Profile", level=1)
    for line in context.splitlines():
        if line.strip():
            doc.add_paragraph(line, style="List Bullet")

    numeric_cols = df.select_dtypes("number").columns.tolist()
    if numeric_cols:
        doc.add_heading("Numeric Column Summary", level=1)
        for col in numeric_cols:
            s = df[col].dropna()
            doc.add_paragraph(
                f"{col}: mean={s.mean():.4g}, std={s.std():.4g}, "
                f"min={s.min():.4g}, max={s.max():.4g}, "
                f"nulls={df[col].isnull().sum()}",
                style="List Bullet",
            )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()